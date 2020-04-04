from docx import Document

from tkinter_message import show_error

regex = r"\{\{([A-Za-z0-9_]+)\}\}"


class MergeDocument:
    def __init__(self, matches: dict, document, path):
        self.document = document
        self.matches = matches
        self.save_path = path

    def replace_matches_newpage(self, max_count):
        para_list = self.document.paragraphs
        for i in range(max_count - 1):
            self.duplicate_page(para_list)
        wait_page_breaks = 1
        for items in self.matches:
            page_breaks_waited = 0
            for paragraph in self.document.paragraphs:
                for key, value in items.items():
                    if ("{{" + key + "}}") in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{" + key + "}}", str(value))
                for run in paragraph.runs:
                    if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                        page_breaks_waited += 1

                if page_breaks_waited >= wait_page_breaks:
                    wait_page_breaks += 1
                    break

    def replace_matches(self):
        for key, value in self.matches.items():
            for paragraph in self.document.paragraphs:
                if ("{{" + key + "}}") in paragraph.text:
                    paragraph.text = paragraph.text.replace("{{" + key + "}}", str(value))
                    print(key)

    def duplicate_page(self, para_list):
        self.document.add_page_break()
        for paragraph in para_list:
            self.get_para_data(paragraph)

    def get_para_data(self, paragraph):

        output_para = self.document.add_paragraph()
        for run in paragraph.runs:
            output_run = output_para.add_run(run.text)
            output_run.bold = run.bold
            output_run.italic = run.italic
            output_run.underline = run.underline
            output_run.font.color.rgb = run.font.color.rgb
            output_run.style.name = run.style.name
        output_para.style = paragraph.style
        output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment




    def save(self):
        print(self.save_path)
        self.document.save(self.save_path)


def init_document(path):
    try:
        return Document(path)
    except FileNotFoundError:
        show_error("File not found", "File at " + str(path) + "not found")







